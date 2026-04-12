from docx import Document
path='/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx'
doc=Document(path)

def has_drawing(p):
    return bool(p._p.xpath('.//w:drawing'))

caps = [
    'Figure 1. Official final benchmark comparison',
    'Figure 2. Flash-crowd case study',
    'Figure 3. Official benchmark summary',
    'Figure 4. Simulator topology playback',
    'Appendix Figure A1. Official locked validation bundles',
    'Appendix Figure A2. Scenario-by-scenario PPO-based hybrid controller',
]
print('paragraphs=', len(doc.paragraphs))
print('inline_shapes=', len(doc.inline_shapes))
print('references_present=', any(p.text.strip()=='References' for p in doc.paragraphs))
print('visible_body_citation_paragraphs=', sum(1 for p in doc.paragraphs if '(' in p.text and ')' in p.text and p.text.strip()!='References'))
for cap in caps:
    matches=[i for i,p in enumerate(doc.paragraphs) if p.text.strip().startswith(cap)]
    print(cap, matches)
    if matches:
        i=matches[0]
        print(' prev_has_drawing=', i>0 and has_drawing(doc.paragraphs[i-1]))
        print(' next_text=', doc.paragraphs[i+1].text[:120] if i+1 < len(doc.paragraphs) else '')

for start,end in [(48,63),(144,163),(179,198)]:
    print('\nRANGE', start, end)
    for i in range(start, min(end, len(doc.paragraphs))):
        p=doc.paragraphs[i]
        print(f'{i:03d} [{p.style.name}] {p.text}')
