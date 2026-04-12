from docx import Document

path='/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx'

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

# pass 1: remove duplicated Valadarsky entry if present
Doc = Document(path)
for i in range(len(Doc.paragraphs)-1):
    a = Doc.paragraphs[i].text.strip()
    b = Doc.paragraphs[i+1].text.strip()
    if a and a == b and 'Valadarsky' in a:
        delete_paragraph(Doc.paragraphs[i+1])
        break
Doc.save(path)

# pass 2: normalize appendix tail
Doc = Document(path)
# find appendix heading
start = None
for i,p in enumerate(Doc.paragraphs):
    if p.text.strip() == 'Appendix: Supporting Figures':
        start = i
        break
if start is None:
    raise SystemExit('appendix heading not found')
texts = {
    start+1: 'The appendix retains real but non-central evidence that supports reproducibility, traceability, and broader interpretation without overloading the main thesis narrative.',
    start+2: 'Appendix Figure A1 documents the locked validation bundles used to store scenario-specific acceptance results under fixed constraints.',
    start+3: 'Appendix Figure A1. Official locked validation bundles used to document scenario-specific acceptance results and reproducibility settings.',
    start+4: 'Appendix Figure A1 is useful because it preserves the chain of evidence behind GreenNet\'s evaluation workflow. It shows that important held-out validations were stored as traceable bundles with explicit seeds, episode counts, and controller settings. However, because these bundles reflect a narrower validation context than the final benchmark, they are treated as supporting evidence rather than as central thesis results.',
    start+5: 'Appendix Figure A2 provides a compact scenario-by-scenario view of how the PPO-based hybrid controller changes energy use and delivered traffic relative to the All-On controller in the official final benchmark.',
    start+6: 'Appendix Figure A2. Scenario-by-scenario PPO-based hybrid controller trade-off relative to the All-On controller in the official final benchmark.',
    start+7: 'Appendix Figure A2 is intentionally placed outside the main body because it is analytically useful but more detailed than necessary for the primary thesis narrative. It helps the reader see that the PPO-based hybrid controller\'s strongest scenario-level case is flash-crowd, that diurnal and hotspot are better interpreted as compromise-policy cases, and that the normal scenario remains a negative result.',
}
for idx, text in texts.items():
    if idx < len(Doc.paragraphs):
        Doc.paragraphs[idx].text = text
        Doc.paragraphs[idx].style = 'Normal'
# blank any extra residual appendix paras before end if they exist nearby
for idx in range(start+8, min(start+12, len(Doc.paragraphs))):
    if Doc.paragraphs[idx].style.name == 'Normal':
        if 'Appendix Figure A' in Doc.paragraphs[idx].text or not Doc.paragraphs[idx].text.strip():
            Doc.paragraphs[idx].text = ''
Doc.save(path)
