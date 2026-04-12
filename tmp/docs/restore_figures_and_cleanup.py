from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Inches

TARGET = Path('/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx')
SOURCE = Path('/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_edited.docx')
BACKUP = Path('/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized_backup_before_figure_restore.docx')

src = Document(str(SOURCE))
doc = Document(str(TARGET))

# Helpers

def get_blobs_by_caption(document):
    out = {}
    for i, p in enumerate(document.paragraphs):
        text = p.text.strip()
        if text.startswith('Figure 4 is retained'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Figure 4'] = document.part.related_parts[rid[0]].blob
        elif text.startswith('Figure 1 is the central'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Figure 1'] = document.part.related_parts[rid[0]].blob
        elif text.startswith('Figure 2 focuses') or text.startswith('Figure 2 narrows'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Figure 2'] = document.part.related_parts[rid[0]].blob
        elif text.startswith('Figure 3 is retained') or text.startswith('Figure 3 serves'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Figure 3'] = document.part.related_parts[rid[0]].blob
        elif text.startswith('Appendix Figure A1 documents'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Appendix Figure A1'] = document.part.related_parts[rid[0]].blob
        elif text.startswith('Appendix Figure A2 provides'):
            draw_p = document.paragraphs[i+1]
            rid = draw_p._p.xpath('.//a:blip/@r:embed')
            if rid:
                out['Appendix Figure A2'] = document.part.related_parts[rid[0]].blob
    return out

blobs = get_blobs_by_caption(src)

tmp_dir = Path('/Users/oltazagraxha/Desktop/GreenNet/tmp/docs/restored_media')
tmp_dir.mkdir(parents=True, exist_ok=True)
media_paths = {}
for name, blob in blobs.items():
    p = tmp_dir / (name.replace(' ', '_').replace(':', '') + '.png')
    p.write_bytes(blob)
    media_paths[name] = p

# Fix hypothesis section and duplicates using current paragraph indices known from inspection.
# 59 label, 60+61 duplicate explanatory paras. Replace 60 with actual statement, keep 61 as explanation with normalized wording.
doc.paragraphs[60].text = (
    'The working hypothesis of this thesis is that the PPO-based hybrid controller can achieve modest energy savings relative to the All-On controller while remaining within explicit QoS limits, but that those gains may not surpass the heuristic controller across the full benchmark.'
)
doc.paragraphs[61].text = (
    'In GreenNet, this expectation is deliberately narrower than a claim that reinforcement learning will outperform every non-AI baseline. The hypothesis is designed to test whether adaptive control can produce a defensible energy-QoS trade-off under controlled conditions, even if the effect is modest and scenario-dependent.'
)

# Remove duplicate paragraphs by deleting their XML nodes.
for idx in [77, 152, 169, 173]:
    p = doc.paragraphs[idx]._p
    p.getparent().remove(p)

# After deletions, reload to refresh indices before figure insertion.
doc.save(str(TARGET))
doc = Document(str(TARGET))

# Insert missing figure image paragraphs before captions if missing.
def para_has_drawing(p):
    return bool(p._p.xpath('.//w:drawing'))

caption_to_name = {
    'Figure 1. Official final benchmark comparison': 'Figure 1',
    'Appendix Figure A1. Official locked validation bundles': 'Appendix Figure A1',
    'Appendix Figure A2. Scenario-by-scenario PPO-based hybrid controller': 'Appendix Figure A2',
}

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for prefix, fig_name in caption_to_name.items():
        if text.startswith(prefix):
            prev = doc.paragraphs[i-1] if i > 0 else None
            if prev is None or not para_has_drawing(prev):
                new_p = p.insert_paragraph_before('')
                run = new_p.add_run()
                run.add_picture(str(media_paths[fig_name]), width=Inches(5.8))
                # center the inserted figure paragraph to fit thesis layout
                new_p.style = p.style
            break

# Normalize a couple of legacy terminology occurrences left in body if any.
for p in doc.paragraphs:
    if 'PPO-based controller' in p.text:
        p.text = p.text.replace('PPO-based controller', 'PPO-based hybrid controller')
    if 'PPO, and Heuristic policies' in p.text:
        p.text = p.text.replace('PPO, and Heuristic policies', 'the heuristic controller, and the PPO-based hybrid controller')
    if p.text.strip() == '':
        continue

# Save
if BACKUP.exists():
    pass

doc.save(str(TARGET))

# Verification summary
final = Document(str(TARGET))
fig_caps = [
    'Figure 1.', 'Figure 2.', 'Figure 3.', 'Figure 4.',
    'Appendix Figure A1.', 'Appendix Figure A2.'
]
body_citations = sum(1 for p in final.paragraphs if '(' in p.text and ')' in p.text and 'References' not in p.text)
print('paragraphs', len(final.paragraphs))
print('inline_shapes', len(final.inline_shapes))
for cap in fig_caps:
    found = [i for i,p in enumerate(final.paragraphs) if p.text.strip().startswith(cap)]
    print(cap, found)
    if found:
        idx = found[0]
        print('  prev_has_drawing', idx>0 and para_has_drawing(final.paragraphs[idx-1]))
print('references_present', any(p.text.strip()=='References' for p in final.paragraphs))
print('body_citation_paragraphs', body_citations)
