from docx import Document


SOURCE = "/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_edited.docx"
TARGET = "/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx"


def move_block_before(doc, start, end, before_idx):
    paragraphs = doc.paragraphs
    block = [paragraphs[i]._p for i in range(start, end + 1)]
    target = paragraphs[before_idx]._p
    for element in block:
        target.addprevious(element)


def move_block_to_end(doc, start, end):
    paragraphs = doc.paragraphs
    block = [paragraphs[i]._p for i in range(start, end + 1)]
    body = doc._body._body
    sect_pr = body.sectPr
    for element in block:
        sect_pr.addprevious(element)


def insert_paragraph_before(paragraph, text, style):
    new_paragraph = paragraph.insert_paragraph_before(text)
    new_paragraph.style = style
    return new_paragraph


doc = Document(SOURCE)

# Update transition into the results chapter so the reordered structure reads cleanly.
doc.paragraphs[141].text = (
    "This section presents the primary results used for the final comparative analysis of "
    "GreenNet. The main body keeps the official final benchmark comparison, one selected "
    "PPO-positive case study, and the acceptance-threshold evidence that supports the "
    "interpretation of the benchmark outcome. More detailed supporting figures are moved to "
    "the appendix so that the central argument remains selective and academically coherent."
)

# Keep validation inside the results chapter but make the link to the chapter explicit.
doc.paragraphs[152].text = (
    "Alongside comparative results, GreenNet also needs to demonstrate that its evaluation "
    "workflow is traceable and criteria-driven. For this reason, the thesis retains the "
    "acceptance-threshold evidence within the results chapter and moves narrower validation "
    "artifacts to the appendix."
)

# Convert the later duplicate implementation heading into a supporting subsection that
# belongs to the main implementation chapter.
doc.paragraphs[157].text = "Implementation Figure: Simulator Topology Playback"
doc.paragraphs[157].style = "Heading 2"
doc.paragraphs[158].text = (
    "Within the implementation chapter, one visual figure is retained. The simulator "
    "topology playback is clearer and easier to defend than a crowded dashboard overview "
    "because it connects system design to observable execution without forcing the thesis "
    "to justify too many run-specific widgets at once."
)

# Reorder the late implementation figure block so it sits under the main implementation chapter.
move_block_before(doc, 157, 163, 103)

# Append the supporting-figures appendix to the end of the thesis, after the conclusion.
move_block_to_end(doc, 164, 173)

# Promote the appendix heading so it clearly reads as an end matter section.
appendix_heading = None
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "Appendix: Supporting Figures":
        paragraph.style = "Heading 1"
        appendix_heading = paragraph
        break

# Add a properly placed references heading near the end without inventing bibliography entries.
if appendix_heading is not None:
    insert_paragraph_before(appendix_heading, "References", "Heading 1")

doc.save(TARGET)
