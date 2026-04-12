from docx import Document
path='/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx'
doc=Document(path)

# Remove adjacent duplicate non-empty paragraphs anywhere in the body.
removed = 0
idx = 1
while idx < len(doc.paragraphs):
    prev = doc.paragraphs[idx-1].text.strip()
    cur = doc.paragraphs[idx].text.strip()
    if prev and cur and prev == cur:
        p = doc.paragraphs[idx]._p
        p.getparent().remove(p)
        removed += 1
    else:
        idx += 1

doc.save(path)

# Verify critical headings/figures after save.
doc=Document(path)
for i,p in enumerate(doc.paragraphs):
    if p.text.strip() in ['Results and Discussion','Conclusion','References','Appendix: Supporting Figures']:
        print(i, p.style.name, p.text.strip())
print('removed', removed)
print('inline_shapes', len(doc.inline_shapes))
