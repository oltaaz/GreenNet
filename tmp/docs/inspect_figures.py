from copy import deepcopy
from docx import Document
from pathlib import Path

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

def para_has_drawing(p):
    return bool(p._p.xpath('.//w:drawing'))

def image_blips(p):
    return [el.get('{%s}embed' % NS['r']) for el in p._p.xpath('.//a:blip')]

for path in [
    '/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx',
    '/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_edited.docx',
]:
    doc = Document(path)
    print('\nFILE', Path(path).name, 'paragraphs', len(doc.paragraphs), 'inline_shapes', len(doc.inline_shapes))
    for i,p in enumerate(doc.paragraphs):
        text = p.text.strip().replace('\n',' ')
        if para_has_drawing(p) or 'Figure' in text or p.style.name.startswith('Heading'):
            marker = 'DRAW' if para_has_drawing(p) else '    '
            print(f'{i:03d} {marker} [{p.style.name}] {text[:140]}')
            if para_has_drawing(p):
                print('    blips=', image_blips(p))
