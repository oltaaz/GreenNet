from docx import Document
path='/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx'
doc=Document(path)
for start,end in [(48,63),(73,79),(147,163),(164,176),(187,205)]:
    print('\nRANGE',start,end)
    for i in range(start,end):
        if i>=len(doc.paragraphs): break
        p=doc.paragraphs[i]
        flag='DRAW' if bool(p._p.xpath('.//w:drawing')) else '    '
        print(f'{i:03d} {flag} [{p.style.name}] {p.text}')
print('\ncounts', len(doc.paragraphs), len(doc.inline_shapes))
