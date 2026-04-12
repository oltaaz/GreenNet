from docx import Document


PATH = "/Users/oltazagraxha/Desktop/GreenNet/output/doc/RIT_Digital_Institutional_GreenNet_reorganized.docx"


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def find_first(doc, startswith):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i, p
    raise ValueError(f"Not found: {startswith}")


doc = Document(PATH)

# Remove misplaced duplicate results paragraph from the Technologies section.
idx, para = find_first(doc, "The main conclusion from Figure 1 is conservative.")
if idx < 130:
    delete_paragraph(para)

# Remove misplaced duplicate results-interpretation paragraph from Challenges.
idx, para = find_first(doc, "Taken together, Figures 1 and 2 support a precise interpretation.")
if idx < 140:
    delete_paragraph(para)

# Replace the misplaced conclusion paragraph in Evaluation Strategy with the correct controlled-comparison paragraph.
idx, para = find_first(doc, "The working expectation for GreenNet was that adaptive control would reduce energy use relative to the All-On controller")
if idx < 150:
    para.text = (
        "Controlled comparison is the first principle. The All-On controller, the heuristic controller, "
        "and the PPO-based hybrid controller are executed under the same topology assumptions, traffic "
        "scenarios, step limits, and metric definitions so that differences in outcome can be traced "
        "back to controller behavior rather than to changing experimental conditions."
    )

# Restore the Conclusion heading that was displaced by a stray appendix caption line.
idx, para = find_first(doc, "Appendix Figure A2. Scenario-by-scenario PPO-based hybrid controller trade-off relative to the All-On controller")
if idx < 170:
    para.text = "Conclusion"
    para.style = "Heading 1"

# Remove the stray appendix A2 lead-in that was inserted before Conclusion.
idx, para = find_first(doc, "Appendix Figure A2 provides a compact scenario-by-scenario view of how the PPO-based hybrid controller changes energy use")
if idx < 170:
    delete_paragraph(para)

# Keep the appendix wording consistent with the thesis terminology.
for p in doc.paragraphs:
    text = p.text.strip()
    if text == "Appendix Figure A2 provides a compact scenario-by-scenario view of how PPO changes energy use and delivered traffic relative to All-On in the official final benchmark.":
        p.text = "Appendix Figure A2 provides a compact scenario-by-scenario view of how the PPO-based hybrid controller changes energy use and delivered traffic relative to the All-On controller in the official final benchmark."
    elif text == "Appendix Figure A2. Scenario-by-scenario PPO trade-off relative to All-On in the official final benchmark.":
        p.text = "Appendix Figure A2. Scenario-by-scenario PPO-based hybrid controller trade-off relative to the All-On controller in the official final benchmark."

doc.save(PATH)
